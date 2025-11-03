"""Convert meeting summary and minutes JSON files into a DOCX report."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document as DocumentFactory
from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


def load_json(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def ensure_custom_styles(document: Document) -> None:
    styles = document.styles

    if "Heading 2" in styles:
        heading = styles["Heading 2"]
        heading.font.size = Pt(16)  # type: ignore[attr-defined]

    if "Heading 3" in styles:
        heading = styles["Heading 3"]
        heading.font.size = Pt(14)  # type: ignore[attr-defined]

    if "Body Text" not in styles:
        body_style = styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.size = Pt(11)  # type: ignore[attr-defined]
        body_style.font.name = "Calibri"  # type: ignore[attr-defined]


def add_heading(document: Document, text: str, level: int = 2) -> None:
    document.add_heading(text, level=level)


def resolve_style(document: Document, name: str, fallback: str = "Normal") -> Any:
    styles = document.styles
    try:
        return styles[name]
    except KeyError:
        return styles[fallback]


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.style = resolve_style(document, "Body Text")


def add_bullet_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        if not item:
            continue
        paragraph = document.add_paragraph(str(item), style="List Bullet")
        paragraph.style.font.size = Pt(11)  # type: ignore[attr-defined]


def add_numbered_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        if not item:
            continue
        paragraph = document.add_paragraph(str(item), style="List Number")
        paragraph.style.font.size = Pt(11)  # type: ignore[attr-defined]


def render_summary(document: Document, summary_json: Dict[str, Any]) -> None:
    meeting_id = summary_json.get("meeting_id", "Unknown Meeting")
    add_heading(document, f"Meeting Summary – {meeting_id}", level=1)

    summaries = summary_json.get("summaries", {})
    exec_summary = summaries.get("executive_summary")
    if exec_summary:
        add_heading(document, "Executive Summary", level=2)
        add_paragraph(document, exec_summary)

    detailed_summary = summaries.get("detailed_summary")
    if detailed_summary:
        add_heading(document, "Detailed Summary", level=2)
        add_paragraph(document, detailed_summary)

    highlights = summaries.get("key_highlights", [])
    if highlights:
        add_heading(document, "Key Highlights", level=2)
        add_bullet_list(document, highlights)

    topics = summaries.get("topics_covered", [])
    if topics:
        add_heading(document, "Topics Covered", level=2)
        for topic in topics:
            if isinstance(topic, dict):
                title = topic.get("title") or topic.get("topic")
                notes = topic.get("notes") or topic.get("details")
                add_paragraph(document, f"• {title}" if title else "• Topic")
                if notes:
                    add_paragraph(document, f"  {notes}")
            else:
                add_paragraph(document, f"• {topic}")

    sentiment = summaries.get("sentiment_analysis")
    if sentiment:
        add_heading(document, "Sentiment Analysis", level=2)
        if isinstance(sentiment, dict):
            for key, value in sentiment.items():
                add_paragraph(document, f"{key.title()}: {value}")
        else:
            add_paragraph(document, str(sentiment))


def render_minutes(document: Document, minutes_json: Dict[str, Any]) -> None:
    minutes = minutes_json.get("minutes", {})
    add_heading(document, "Meeting Minutes", level=1)

    meeting_info = minutes.get("meeting_info", {})
    if meeting_info:
        add_heading(document, "Meeting Information", level=2)
        table = document.add_table(rows=0, cols=2)
        table.style = "Light List Accent 1"
        for key in ("title", "date", "start_time", "end_time", "attendees"):
            value = meeting_info.get(key)
            if not value:
                continue
            row_cells = table.add_row().cells
            row_cells[0].text = key.replace("_", " ").title()
            if isinstance(value, list):
                row_cells[1].text = ", ".join(str(item) for item in value)
            else:
                row_cells[1].text = str(value)
        document.add_paragraph("")

    agenda_items = minutes.get("agenda_items", [])
    if agenda_items:
        add_heading(document, "Agenda Items", level=2)
        for agenda in agenda_items:
            title = agenda.get("topic", "Agenda Item")
            add_heading(document, title, level=3)
            if agenda.get("summary"):
                add_paragraph(document, agenda["summary"])
            if agenda.get("discussion_points"):
                add_heading(document, "Discussion Points", level=4)
                add_bullet_list(document, agenda["discussion_points"])
            if agenda.get("decisions"):
                add_heading(document, "Decisions", level=4)
                add_numbered_list(document, agenda["decisions"])

    action_items = minutes.get("action_items", [])
    if action_items:
        add_heading(document, "Action Items", level=2)
        for item in action_items:
            header = f"Action {item.get('id', '')}: {item.get('description', 'Action Item')}"
            add_heading(document, header.strip(), level=3)
            details = [
                ("Owner", item.get("owner")),
                ("Due Date", item.get("due_date")),
                ("Status", item.get("status")),
                ("Priority", item.get("priority")),
            ]
            for label, value in details:
                if value:
                    add_paragraph(document, f"{label}: {value}")

    announcements = minutes.get("announcements", [])
    if announcements:
        add_heading(document, "Announcements", level=2)
        add_bullet_list(document, announcements)

    next_meeting = minutes.get("next_meeting")
    if next_meeting:
        add_heading(document, "Next Meeting", level=2)
        if isinstance(next_meeting, dict):
            for key, value in next_meeting.items():
                if value:
                    if isinstance(value, list):
                        add_paragraph(document, f"{key.title()}: {', '.join(map(str, value))}")
                    else:
                        add_paragraph(document, f"{key.title()}: {value}")
        else:
            add_paragraph(document, str(next_meeting))


def build_document(summary_path: Path, minutes_path: Path, output_path: Path) -> None:
    summary_data = load_json(summary_path)
    minutes_data = load_json(minutes_path)

    document = DocumentFactory()
    ensure_custom_styles(document)

    render_summary(document, summary_data)

    document.add_page_break()

    render_minutes(document, minutes_data)

    document.save(str(output_path))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("summary_json", type=Path, help="Path to 04_summaries.json")
    parser.add_argument("minutes_json", type=Path, help="Path to 03_minutes.json")
    parser.add_argument(
        "output_docx",
        type=Path,
        help="Output DOCX filename (e.g., meeting_report.docx)",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    build_document(args.summary_json, args.minutes_json, args.output_docx)


if __name__ == "__main__":
    main()
